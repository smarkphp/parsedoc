'''
Desccription: 
version: 
company: zy
Author: Mark
Date: 2020-10-02 10:12:29
ListenEditors: Mark
LastEditTime: 2020-10-07 18:08:07
'''

from docx import Document
from docx.shared import Inches
import os
import json
# reload(sys)
# sys.setdefaultencoding('utf-8')

class parse:


    @staticmethod
    def sum(a, b):
        c = a + b
        print(c)
   
    '''写docx文档代码'''
    # @classmethod 
    @staticmethod
    def write_doc(a):
        print(a)
       
     
        document = Document()
        document.add_heading(u'我的一个新文档',0)

        p = document.add_paragraph('A plain paragraph having some ')   #插入段落
        p.add_run('bold').bold = True
        p.add_run(' and some ')
        p.add_run('italic.').italic = True

        # 添加一级标题
        document.add_heading('Heading, level 1', level=1)
        document.add_paragraph('Intense quote', style='IntenseQuote')

        # 添加无序列表
        document.add_paragraph('first item in unordered list', style='ListBullet' )

        # 添加有序列表
        document.add_paragraph('first item in ordered list',style = 'ListNumber')
        document.add_paragraph('second item in ordered list',style = 'ListNumber')
        document.add_paragraph('third item in ordered list',style = 'ListNumber')

        # 插入图片
        # document.add_picture('monty-truth.png', width=Inches(1.25)) #插入图片

        # # >> 添加表格1
        table = document.add_table(rows = 1,cols = 3)
        # 获取第一行的单元格列表对象
        hdr_cells = table.rows[0].cells
        # 为每一个单元格赋值
        # 注：值都要为字符串类型
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Age'
        hdr_cells[2].text = 'Tel'
        # # 为表格添加一行
        new_cells = table.add_row().cells
        new_cells[0].text = 'Tom'
        new_cells[1].text = '19'
        new_cells[2].text = '12345678'

        new_cells = table.add_row().cells
        new_cells[0].text = 'Alice'
        new_cells[1].text = '20'
        new_cells[2].text = '73926589'
        
        document.add_paragraph('second item in unordered list', style='ListBullet' )
        # # >> 添加表格2
        table2 = document.add_table(rows = 1,cols = 2)
        hdr_cells1 = table2.rows[0].cells
        hdr_cells1[0].text = 'Name'
        hdr_cells1[1].text = 'Age'

        new_cells1 = table2.add_row().cells
        new_cells1[0].text = 'Tom1'
        new_cells1[1].text = '19'

        new_cells1 = table2.add_row().cells
        new_cells1[0].text = 'Alice1'
        new_cells1[1].text = '20'

        # # 添加分页符
        document.add_page_break()
        
        # # 往新的一页中添加段落
        p = document.add_paragraph('This is a paragraph in new page.')
        p = document.add_paragraph('内容是指事物所包含的实质性事物。一个艺术作品的表现、基本含义、意味或审美价值。内容关涉的是我们在艺术作品中感受到的感觉的、主观的、心理的和情感方面的意涵，与我们对描绘性方面的单纯知觉相对')
        # 保存文档
        document.save('demo1.docx')

    '''读取文档内容: 文件、关键字'''
    @staticmethod
    def read_doc(file):
        print(">>读取文档内容")
        # 创建文档对象
        document = Document(file)

        # 读取文档中所有的段落列表
        ps = document.paragraphs

        for x in ps:
            if '实质性事物' in x.text:
                print(x.text)
        print("----\n----")
        # 每个段落有两个属性：style和text
        ps_detail = [(x.text,x.style.name) for x in ps]
        # print(ps_detail,type(ps_detail))
        for pd in ps_detail:
            print("pd:",pd,type(pd))
        


        with open('out.tmp','w+') as fout:
            fout.write('')  
        # 读取段落并写入一个文件
        with open('out.tmp','a+') as fout:
            for p in ps_detail:
                fout.write(p[0] + '\t' + p[1] + '\n\n')

        # 读取文档中的所有段落的列表
        tables = document.tables
        with open('out.tmp','a+') as fout:
            for table in tables:
                for row in table.rows:
                    for cell in row.cells:
                        #  print(cell.text)
                        #  fout.write(cell.text + '\t')
                         if cell.text == "Tom":
                             print(cell.text)
                    fout.write('\n')

        os.system('rm -rf out.tmp')

    @staticmethod 
    def read_doc_common(file, keyword):
        document = Document(file)
        from pprint import pprint
 
        dictV={}
        lrv =[]
        # 读取文档中所有的段落列表-{文字}
        ps = document.paragraphs
        for x in ps:
            if keyword in x.text:
                # print("graph:",x.text)
                lrv.append(x.text)

        print("----\n----")


        # 读取文档中的所有段落的列表-{表格}
        tables = document.tables
        for table in tables:
            for row in table.rows:
                # print("rc: ",row.text)
                for cell in row.cells:
                        if keyword in cell.text:
                            # print("table1: ",cell.text)
                            lrv.append(cell.text)
        # print("lrv:",lrv)
        
        # TODO 根据具体情况可能需要写到单一个column对应到其它column
        # TODO for i in range(1, len(table.rows)):#从表格第二行开始循环读取表格数据：类似于如下
        ## idNum = table.cell(i,0).text #序号
        ## companyName = table.cell(i,1).text  #控股企业名称
        ## investmentRate = table.cell(i,2).text   #投资比例
        ## stock= table.cell(i,3).text  #股权链

        # lrvJson = json.dumps(lrv,ensure_ascii=False)
        # print(type(lrv))
        # print(type(lrvJson))
        
        dictV[keyword]=lrv
        return dictV
        

if __name__ == "__main__":
    #    parse.sum(5,6)
    # parse.write_doc(5)

    # lrvd = parse.read_doc_common('../file/doc/项目管理样例eg.docx','项目实施')
    # print(lrvd)
    # print('\n',lrvd['项目实施'][2])

    # i = 1
    # for lv in lrvd['项目实施']:
    #     print(i,':',lv)
    #     i=i+1

    # print("\n")
    dpt = lrvjson = parse.read_doc_common('../file/doc/demo1.docx','Tom')
    print(dpt)
    
    # print("lrvjson:",lrvjson)
    # print("lrvjson-dt:",lrvjson['Tom'],',type: ',type(lrvjson['Tom']))